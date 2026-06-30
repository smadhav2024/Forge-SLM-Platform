import os
import re
import fitz  # PyMuPDF
import pymupdf4llm
from fastembed import TextEmbedding
from docx import Document as DocxDocument
from app.database import AsyncSessionLocal
from app.models import DocumentVector

# Initialize the embedding model globally
embedding_model = TextEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Matches a full GitHub-flavored markdown table block (header + separator + rows).
# The optional tail alternative also matches a final row with no trailing
# newline, since chunk_markdown's table chunks often end without one.
_TABLE_BLOCK_RE = re.compile(
    r"(?:^\|.*\|[ \t]*\n)+(?:^\|.*\|[ \t]*$)?|^\|.*\|[ \t]*$",
    re.MULTILINE,
)
_TABLE_SEPARATOR_RE = re.compile(r"^\|[\s:|-]+\|[ \t]*$", re.MULTILINE)
_HEADING_RE = re.compile(r"^(#{1,6})\s+.*$", re.MULTILINE)


def _table_block_to_sentences(table_text: str) -> str:
    """Converts a markdown table block into plain "key: value, key: value"
    sentences per row, using the header row as field names.

    Embedding models like all-MiniLM-L6-v2 are trained on natural language;
    raw pipe/dash table syntax (|, ---, :--) adds symbol noise that dilutes
    similarity against a plain-language question. This keeps the same
    information but phrased as text the embedding space actually represents.
    """
    rows = [r.strip() for r in table_text.strip().split("\n") if r.strip()]
    rows = [r for r in rows if not _TABLE_SEPARATOR_RE.match(r)]
    if not rows:
        return ""

    def cells(row):
        return [c.strip() for c in row.strip("|").split("|")]

    header = cells(rows[0])
    sentences = []
    for row in rows[1:]:
        values = cells(row)
        pairs = [
            f"{h}: {v}" for h, v in zip(header, values) if h and v
        ]
        if pairs:
            sentences.append(", ".join(pairs) + ".")
    return " ".join(sentences)


def _to_embedding_text(chunk: str) -> str:
    """Strips markdown syntax from a chunk for embedding purposes only.

    The chunk stored in the DB (and later sent to the LLM as context) stays
    markdown-formatted, since the LLM reads markdown fine. But the same
    chunk fed into the embedding model was previously polluting the vector
    with structural symbols (#, |, ---), which moved it away from how a
    plain-language question embeds, hurting retrieval. This produces a
    clean-text version used only as the embedding input.
    """
    text = chunk

    # Convert table blocks to plain sentences before stripping pipes generally.
    text = _TABLE_BLOCK_RE.sub(lambda m: _table_block_to_sentences(m.group(0)) + "\n", text)

    # Strip heading markers but keep the heading text (it's useful context).
    text = _HEADING_RE.sub(lambda m: m.group(0).lstrip("#").strip(), text)

    # Strip residual markdown emphasis/code markers that don't carry meaning.
    text = re.sub(r"[*_`]{1,3}", "", text)

    # Collapse excess whitespace left behind by the substitutions.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _split_into_sections(markdown_text: str):
    """Splits markdown into sections at heading boundaries.

    Each section keeps its heading line attached to the content that follows,
    so chunks retain their surrounding context instead of being sliced blindly.
    """
    matches = list(_HEADING_RE.finditer(markdown_text))
    if not matches:
        return [markdown_text]

    sections = []
    if matches[0].start() > 0:
        sections.append(markdown_text[: matches[0].start()])

    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_text)
        sections.append(markdown_text[start:end])

    return [s for s in sections if s.strip()]


def _split_section_keeping_tables_whole(section: str):
    """Splits a section into ordered blocks of plain text and whole markdown tables.

    Tables are never sliced mid-row; each table block is kept as a single
    contiguous unit so chunking can treat it atomically.
    """
    blocks = []
    last_end = 0
    for match in _TABLE_BLOCK_RE.finditer(section):
        if match.start() > last_end:
            blocks.append(("text", section[last_end:match.start()]))
        blocks.append(("table", match.group(0)))
        last_end = match.end()
    if last_end < len(section):
        blocks.append(("text", section[last_end:]))
    return blocks


def _chunk_plain_text(text: str, chunk_size: int, overlap: int):
    """Size-bounded chunking for prose, splitting on paragraph/sentence boundaries
    where possible instead of cutting at arbitrary character offsets."""
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    # Prefer splitting on paragraph breaks, then sentence breaks, then whitespace.
    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""

    def flush(buf):
        buf = buf.strip()
        if buf:
            chunks.append(buf)

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(para) > chunk_size:
            # Paragraph itself is too long; fall back to sentence-aware slicing.
            if current:
                flush(current)
                current = ""
            sentences = re.split(r"(?<=[.!?])\s+", para)
            buf = ""
            for sent in sentences:
                if len(buf) + len(sent) + 1 > chunk_size:
                    flush(buf)
                    # carry overlap forward
                    buf = buf[-overlap:] + " " + sent if overlap else sent
                else:
                    buf = (buf + " " + sent).strip()
            if buf:
                current = buf
            continue

        candidate = (current + "\n\n" + para).strip() if current else para
        if len(candidate) > chunk_size:
            flush(current)
            current = (current[-overlap:] + "\n\n" + para).strip() if overlap and current else para
        else:
            current = candidate

    flush(current)
    return chunks


def chunk_markdown(markdown_text: str, chunk_size: int = 500, overlap: int = 50):
    """Splits markdown into reliable, structure-aware chunks.

    - Splits on headings first, so each chunk stays within one logical section.
    - Keeps markdown tables intact as standalone chunks (never split mid-table,
      never interleaved with unrelated prose), since tables collapsed into
      character-offset slices were the main source of unusable chunks.
    - Size-bounds remaining prose using paragraph/sentence-aware slicing with
      overlap, instead of blind fixed-width character windows.
    """
    if overlap >= chunk_size:
        overlap = chunk_size // 2

    chunks = []
    for section in _split_into_sections(markdown_text):
        heading_match = _HEADING_RE.match(section.strip())
        heading_line = heading_match.group(0).strip() if heading_match else None

        blocks = _split_section_keeping_tables_whole(section)
        first_content_seen = False

        for kind, block in blocks:
            if kind == "table":
                table_text = block.strip()
                if not table_text:
                    continue
                # Prepend the section heading to the first table/text block in
                # the section so a heading never becomes its own tiny chunk
                # and context (which section a table belongs to) is preserved.
                prefix = f"{heading_line}\n\n" if heading_line and not first_content_seen else ""
                first_content_seen = True

                # Large tables are still split, but only on row boundaries.
                if len(table_text) <= chunk_size * 3:
                    chunks.append(prefix + table_text)
                else:
                    rows = table_text.split("\n")
                    table_header = "\n".join(rows[:2])  # header + separator row
                    buf = prefix + table_header
                    base = table_header
                    for row in rows[2:]:
                        if len(buf) + len(row) + 1 > chunk_size * 3:
                            chunks.append(buf)
                            buf = base + "\n" + row
                        else:
                            buf += "\n" + row
                    if buf.strip() != base.strip():
                        chunks.append(buf)
            else:
                # Strip the heading line itself out of the text block (it's
                # handled separately above) before chunking the prose.
                text_block = block
                if heading_match and not first_content_seen:
                    text_block = _HEADING_RE.sub("", block, count=1)

                for piece in _chunk_plain_text(text_block, chunk_size, overlap):
                    if heading_line and not first_content_seen:
                        chunks.append(f"{heading_line}\n\n{piece}")
                        first_content_seen = True
                    else:
                        chunks.append(piece)

        # Section had a heading but no body content at all (e.g. trailing
        # heading at end of doc) - still keep it for context, attached
        # rather than dropped.
        if heading_line and not first_content_seen:
            chunks.append(heading_line)

    return [c for c in chunks if c.strip()]


def extract_markdown(file_path: str) -> str:
    """Converts a document to markdown, preserving structure (headings, tables)
    instead of flattening it into raw positional text.

    File type is identified by signature (magic bytes) rather than extension.
    """
    with open(file_path, "rb") as f:
        header = f.read(4)

    if header == b"%PDF":
        # pymupdf4llm walks the page layout (via PyMuPDF) and reconstructs
        # headings, paragraphs, and tables as markdown, instead of the raw
        # token-by-token text dump that page.get_text("text") produces, which
        # is what was leaking table cell text/coordinate-adjacent fragments
        # into chunks.
        markdown_text = pymupdf4llm.to_markdown(file_path)

    elif header.startswith(b"PK"):
        # ZIP signature -> DOCX. Reconstruct a light markdown structure so the
        # same heading/table-aware chunker applies uniformly.
        doc = DocxDocument(file_path)
        lines = []
        for paragraph in doc.paragraphs:
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            text = paragraph.text
            if not text.strip():
                continue
            if style.startswith("heading 1") or style == "title":
                lines.append(f"# {text}")
            elif style.startswith("heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("heading 3"):
                lines.append(f"### {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("")
            rows = table.rows
            if not rows:
                continue
            header_cells = [c.text.strip() for c in rows[0].cells]
            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
            for row in rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        markdown_text = "\n".join(lines)

    else:
        # Fallback to plain text, treated as a single unstructured section.
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            markdown_text = f.read()

    # Clean any stray null bytes to protect PostgreSQL
    return markdown_text.replace("\x00", "")


async def process_document_task(
    conversation_id: int,
    file_path: str,
    source_filename: str = "",
    chunk_size: int = 500,
    overlap: int = 50,
):
    """Background worker that processes the document asynchronously.

    Pipeline: file -> markdown (structure-preserving) -> chunks (heading and
    table aware), instead of the previous file -> raw text -> fixed-width
    character slices, which fragmented tables and produced unusable chunks.
    """
    try:
        markdown_text = extract_markdown(file_path)

        if not markdown_text.strip():
            print(f"SYSTEM WARNING: No text could be extracted from {file_path}")
            return

        text_chunks = chunk_markdown(markdown_text, chunk_size=chunk_size, overlap=overlap)

        if not text_chunks:
            print(f"SYSTEM WARNING: No chunks produced from {file_path}")
            return

        embedding_inputs = [_to_embedding_text(c) for c in text_chunks]
        embeddings = list(embedding_model.embed(embedding_inputs))

        async with AsyncSessionLocal() as db:
            for chunk, emb in zip(text_chunks, embeddings):
                doc_vec = DocumentVector(
                    conversation_id=conversation_id,
                    source_filename=source_filename or None,
                    text_chunk=chunk,
                    embedding_matrix=emb.tolist()
                )
                db.add(doc_vec)

            await db.commit()
            print(
                f"SYSTEM: Successfully embedded {len(text_chunks)} chunks "
                f"(size={chunk_size}, overlap={overlap}) for Conv #{conversation_id}"
            )

    except Exception as e:
        print(f"SYSTEM: RAG Worker Error on Conv #{conversation_id}: {str(e)}")

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)